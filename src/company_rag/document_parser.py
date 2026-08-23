# -*- coding: utf-8 -*-
"""文档解析：PDF / DOCX / TXT / MD → 按页/段落的文本 + 元数据。

- PDF：pypdf 逐页抽取；扫描件（无文本层）检测并提示，可选用 OCR（pytesseract）。
- DOCX：python-docx 抽取段落 + 表格。
- TXT / MD：整篇读取（UTF-8 容错）。
"""
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from .config import settings

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


@dataclass
class ParsedPage:
    page_no: int  # 1-based；非分页文档为 1
    text: str
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    doc_id: str
    filename: str
    source_type: str
    pages: list[ParsedPage]
    metadata: dict = field(default_factory=dict)
    error: Optional[str] = None

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)

    @property
    def page_count(self) -> int:
        return len(self.pages)


def guess_source_type(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    return ext.lstrip(".")


def _clean(text: str) -> str:
    """清洗：规范化空白，去控制字符。"""
    import re

    text = text.replace("\x00", "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_pdf(path: Path, doc_id: str, filename: str) -> ParsedDocument:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[ParsedPage] = []
    n = len(reader.pages)
    # 图片提取（pymupdf，纯 wheel 无系统依赖）
    images_dir = settings.data_dir / "images" / doc_id
    if images_dir.exists():
        import shutil

        shutil.rmtree(images_dir, ignore_errors=True)  # 幂等：先清旧图
    images_dir.mkdir(parents=True, exist_ok=True)
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as e:  # noqa: BLE001
            logger.warning("PDF 第 %d 页抽取失败: %s", i, e)
            text = ""
        text = _clean(text)
        if not text and settings.scan_pdf_ocr:
            text = _ocr_page(path, i)
        page_images = _extract_page_images(path, doc_id, i, images_dir)
        pages.append(ParsedPage(page_no=i, text=text, metadata={"images": page_images}))

    scanned_pages = sum(1 for p in pages if len(p.text) < 20)
    if scanned_pages and n > 0 and scanned_pages / n > 0.5:
        logger.warning("文档 %s 疑似扫描件（%d/%d 页无文本层）；如需检索请启用 OCR", filename, scanned_pages, n)

    return ParsedDocument(
        doc_id=doc_id,
        filename=filename,
        source_type="pdf",
        pages=pages,
        metadata={"page_count": n, "scanned_pages": scanned_pages, "image_dir": str(images_dir)},
    )


def _extract_page_images(path: Path, doc_id: str, page_no: int, images_dir: Path) -> list[str]:
    """提取某页图片，保存到 data/images/{doc_id}/page_{n:03d}/，返回相对文件名列表。"""
    rels: list[str] = []
    try:
        import fitz  # pymupdf
    except ImportError:
        return rels
    try:
        doc = fitz.open(str(path))
        page = doc[page_no - 1]
        imgs = page.get_images(full=True)
        for i, img in enumerate(imgs):
            xref = img[0]
            info = doc.extract_image(xref)
            width, height = info.get("width", 0), info.get("height", 0)
            # 过滤装饰性小图（图标/logo/噪声）
            if width * height < 15000 or len(info["image"]) < 6 * 1024:
                continue
            ext = info["ext"] or "png"
            page_dir = images_dir / f"page_{page_no:03d}"
            page_dir.mkdir(parents=True, exist_ok=True)
            name = f"img_{i}.{ext}"
            (page_dir / name).write_bytes(info["image"])
            rels.append(f"page_{page_no:03d}/{name}")
        doc.close()
    except Exception as e:  # noqa: BLE001
        logger.warning("提取 PDF 图片失败 %s 第%d页: %s", path.name, page_no, e)
    return rels


def _ocr_page(path: Path, page_no: int) -> str:
    """扫描页 OCR（可选依赖 pytesseract + tesseract 可执行文件）。"""
    try:
        import pytesseract  # type: ignore
        from pdf2image import convert_from_path  # type: ignore

        img = convert_from_path(str(path), first_page=page_no, last_page=page_no)[0]
        return pytesseract.image_to_string(img, lang="chi_sim+eng")
    except Exception as e:  # noqa: BLE001
        logger.warning("OCR 不可用: %s（安装 pytesseract/pdf2image 及 tesseract 后启用）", e)
        return ""


def _extract_docx(path: Path, doc_id: str, filename: str) -> ParsedDocument:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    # 正文段落
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    # 表格（每行转文本）
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = _clean("\n".join(parts))
    return ParsedDocument(
        doc_id=doc_id,
        filename=filename,
        source_type="docx",
        pages=[ParsedPage(page_no=1, text=text)],
        metadata={"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)},
    )


def _extract_plain(path: Path, doc_id: str, filename: str, source_type: str) -> ParsedDocument:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            text = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = raw.decode("utf-8", errors="replace")
    text = _clean(text)
    return ParsedDocument(
        doc_id=doc_id,
        filename=filename,
        source_type=source_type,
        pages=[ParsedPage(page_no=1, text=text)],
    )


def parse_file(path: Path, doc_id: Optional[str] = None) -> ParsedDocument:
    """解析单个文件。doc_id 缺省取文件名。"""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持的文件类型 {ext}，支持: {sorted(SUPPORTED_EXTENSIONS)}")

    doc_id = doc_id or path.stem
    try:
        if ext == ".pdf":
            return _extract_pdf(path, doc_id, path.name)
        if ext == ".docx":
            return _extract_docx(path, doc_id, path.name)
        return _extract_plain(path, doc_id, path.name, ext.lstrip("."))
    except Exception as e:  # noqa: BLE001
        logger.exception("解析失败 %s", path)
        return ParsedDocument(
            doc_id=doc_id,
            filename=path.name,
            source_type=ext.lstrip("."),
            pages=[],
            error=str(e),
        )
